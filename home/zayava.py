# -*- coding: utf8 -*-
'''
Установка дополнительных библиотек: pip install python-docx
Версия библиотеки: 0.8.10
Зависимости: Python 2.6, 2.7, 3.3, or 3.4     lxml >= 2.3.2
Документация: https://python-docx.readthedocs.io/en/latest/
'''
from docx import Document
from docx.shared import Pt

def writeTextToDoc(listText, firstRow, listFirstTable, secondRow, listSecondTable):
    document = Document('/home/andrey/Desktop/abitur/home/shablon.docx')

    needStrings = [3,4,5,7,10]
    count = 1
    for paragraphs in document.paragraphs:
        if count == 3:
            paragraphs.text = 'Я, %s %s %s' % (listText[0], listText[1], listText[2])
            for run in paragraphs.runs:
                font = run.font
                font.size = Pt(12)
                font.name = 'Times New Roman'
        if count == 4:
            paragraphs.text = 'Паспорт серия %s № %s Комплектующий орган: %s' % (listText[3], listText[4], listText[5])
            for run in paragraphs.runs:
                font = run.font
                font.size = Pt(12)
                font.name = 'Times New Roman'
        if count == 5:
            paragraphs.text = 'Прошу зарегистрировать меня в качестве кандидата на поступление в Московский университет МВД России имени В.Я. Кикотя в 2020 году на специальность (специализацию): %s, факультет: %s' % (listText[6], listText[7])
            for run in paragraphs.runs:
                font = run.font
                font.size = Pt(12)
                font.name = 'Times New Roman'
        if count == 7:
            paragraphs.text = 'Результаты ЕГЭ:                     Русский язык –    %s    Обществознание – %s                       Математика –     %s      Биология – %s' % (listText[8],listText[9],listText[10],listText[11])
            for run in paragraphs.runs:
                font = run.font
                font.size = Pt(12)
                font.name = 'Times New Roman'
        if count == 10:
            paragraphs.text = 'Мой контактный телефон %s, E-mail: %s' % (listText[12], listText[13])
            for run in paragraphs.runs:
                font = run.font
                font.size = Pt(12)
                font.name = 'Times New Roman'
        count += 1

    for i in range(2,5):
        document.tables[i].rows[0].cells[1].text = '%s\n%s\n%s' % (listText[0], listText[1], listText[2])
        for run in document.tables[i].rows[0].cells[1].paragraphs[0].runs:
                font = run.font
                font.size = Pt(12)
                font.name = 'Times New Roman'


    table = document.tables[0]
    for i in range(firstRow-1):
        table.add_row()

    count = 0
    for row in table.rows:
        for cell in row.cells:
            cell.text = str(listFirstTable[count])
            count += 1
            for run in cell.paragraphs[0].runs:
                font = run.font
                font.size = Pt(12)
                font.name = 'Times New Roman'



    table = document.tables[1]
    for i in range(secondRow-1):
        table.add_row()

    count = 0
    for row in table.rows:
        for cell in row.cells:
            cell.text = str(listSecondTable[count])
            count += 1
            for run in cell.paragraphs[0].runs:
                font = run.font
                font.size = Pt(12)
                font.name = 'Times New Roman'



    # document.save('C:/Users/UzziMauzer/Desktop/AbitAccParser/tmp.docx')
    return document


if __name__ == "__main__":
    listText = ['Сизов',
                'Григорий',
                'Алексеевич',
                '1234',
                '123456',
                'ГУ МВД России по г.Балашиха',
                'Правовое обеспечение национальной безопасности, специализация – уголовно-правовая, предварительное следствие в ОВД',
                'ФПСОИБ',
                '100',
                '95',
                '20',
                'НЕТ',
                '+7(999)123-45-67',
                'tema@tema.ru']

    firstRow = 2
    listFirstTable = ['олимпиада по математике','первое место','олимпиада по японскому','8 место']

    secondRow = 1
    listSecondTable = ['Значек ГТО','Золото']

    writeTextToDoc(listText,firstRow,listFirstTable,secondRow,listSecondTable)
